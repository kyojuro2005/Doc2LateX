from __future__ import annotations

import os
import shutil
import subprocess
import re
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
import docx
import openpyxl
from google import genai
from google.genai import types

from .config import RUNTIME_DIR, GEMINI_API_KEY, GEMINI_MODEL


def tex_template(document_body: str, use_exam_class: bool = False) -> str:
    document_class = 'exam' if use_exam_class else 'article'
    return f"""\\documentclass[12pt]{{{document_class}}}
\\usepackage[T1]{{fontenc}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{amsmath,amssymb,amsfonts,mathtools}}
\\usepackage{{booktabs}}
\\usepackage{{geometry}}
\\geometry{{margin=2.5cm}}
\\usepackage{{hyperref}}
\\usepackage{{enumitem}}
\\usepackage{{graphicx}}
\\usepackage{{array}}
\\begin{{document}}
{document_body}
\\end{{document}}
"""


def escape_latex(text: str) -> str:
    replacements = {
        '\\': r'\textbackslash{}',
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
    }
    return ''.join(replacements.get(ch, ch) for ch in text)


def clean_llm_latex_response(text: str) -> str:
    text = text.strip()
    # Strip markdown codeblocks if present (e.g. ```latex ... ```)
    if text.startswith('```'):
        lines = text.splitlines()
        if lines and lines[0].startswith('```'):
            lines = lines[1:]
        if lines and lines[-1].strip() == '```':
            lines = lines[:-1]
        text = '\n'.join(lines).strip()

    # If the response already contains a full document, extract the body
    if r'\begin{document}' in text and r'\end{document}' in text:
        match = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', text, re.DOTALL)
        if match:
            return match.group(1).strip()
    return text


def call_gemini(contents: list[Any]) -> str:
    if not GEMINI_API_KEY:
        raise RuntimeError("Aucune clé API Gemini configurée.")

    client = genai.Client(api_key=GEMINI_API_KEY)
    models_to_try = [GEMINI_MODEL, 'gemini-3-flash-preview', 'gemini-3.6-flash', 'gemini-3.5-flash', 'gemini-flash-latest']
    seen = set()
    models_to_try = [m for m in models_to_try if m and not (m in seen or seen.add(m))]

    last_error = None
    for model_name in models_to_try:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=contents
            )
            body = response.text or ''
            if body:
                return clean_llm_latex_response(body)
        except Exception as e:
            last_error = e
            continue

    raise RuntimeError(f"Erreur Gemini : {last_error}")


def image_to_latex(image_path: Path) -> str:
    with image_path.open('rb') as f:
        image_bytes = f.read()

    suffix = image_path.suffix.lower()
    media_type_map = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.webp': 'image/webp'}
    media_type = media_type_map.get(suffix, 'image/png')

    prompt = (
        "Vous êtes un expert mondial en typographie LaTeX et en transcription mathématique de documents.\n"
        "Analysez l'image fournie et transcrivez fidèlement et exhaustivement tout le contenu en code LaTeX de haute qualité.\n"
        "Directives rigoureuses :\n"
        "1. Conservez la hiérarchie des titres (\\section, \\subsection, \\subsubsection).\n"
        "2. Formatez TOUTES les équations et formules mathématiques avec une syntaxe rigoureuse ($...$ en ligne, \\[...\\] ou \\begin{equation} en bloc, \\frac, \\sum, matrices, symboles grecs, etc.).\n"
        "3. Convertissez les tableaux en environnements LaTeX 'tabular' ou 'table' avec le package booktabs (\\toprule, \\midrule, \\bottomrule).\n"
        "4. Utilisez des listes ordonnées ou à puces (\\begin{enumerate}, \\begin{itemize}).\n"
        "5. Ne retournez QUE le corps du document LaTeX (sans balises markdown ```latex et sans explications)."
    )

    image_part = types.Part.from_bytes(data=image_bytes, mime_type=media_type)
    clean_body = call_gemini([image_part, prompt])
    return tex_template(clean_body)


def pdf_to_latex(pdf_path: Path) -> str:
    # Try Gemini Multimodal Vision first for optimal rendering of PDF equations, tables, layouts
    if GEMINI_API_KEY:
        try:
            doc = fitz.open(pdf_path)
            parts: list[Any] = []
            max_pages = min(len(doc), 10)  # Render up to 10 pages for multimodal analysis

            for page_num in range(max_pages):
                pix = doc[page_num].get_pixmap(dpi=150)
                img_bytes = pix.tobytes(output='png')
                parts.append(types.Part.from_bytes(data=img_bytes, mime_type='image/png'))

            prompt = (
                f"Vous êtes un expert en conversion de PDF en LaTeX. Voici les {max_pages} pages d'un document PDF rendu en images.\n"
                "Transcrivez l'intégralité du texte, des formules mathématiques ($...$, \\[...\\]), des théorèmes, listes et tableaux en code LaTeX propre et structuré.\n"
                "Ne retournez QUE le corps du document LaTeX sans balises markdown."
            )
            parts.append(prompt)
            clean_body = call_gemini(parts)
            return tex_template(clean_body)
        except Exception as e:
            print(f"Gemini PDF conversion fallback: {e}")

    # Fallback to local text extraction
    doc = fitz.open(pdf_path)
    body: list[str] = []
    for page in doc:
        text = page.get_text('text').strip()
        if text:
            for line in text.splitlines():
                line = line.strip()
                if line:
                    body.append(escape_latex(line) + '\\\\')

    return tex_template('\n'.join(body) if body else '% Document PDF sans texte extractible')


def docx_to_latex(docx_path: Path) -> str:
    document = docx.Document(docx_path)

    # Extract structured text and tables
    extracted_text_parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ''
            extracted_text_parts.append(f"[{style}] {text}" if style else text)

    for table_idx, table in enumerate(document.tables):
        extracted_text_parts.append(f"\n--- TABLEAU {table_idx + 1} ---")
        for row in table.rows:
            row_vals = [cell.text.strip() for cell in row.cells]
            extracted_text_parts.append(" | ".join(row_vals))
        extracted_text_parts.append("--- FIN TABLEAU ---\n")

    full_text = "\n".join(extracted_text_parts)

    # Use Gemini to produce beautiful structured LaTeX from the Word document
    if GEMINI_API_KEY and full_text.strip():
        try:
            prompt = (
                "Vous êtes un expert typographique en LaTeX. Convertissez le contenu suivant extrait d'un document Word (.docx) "
                "en un document LaTeX structuré de haute qualité académique.\n"
                "Directives :\n"
                "1. Transformez les titres en \\section, \\subsection, \\subsubsection.\n"
                "2. Détectez et convertissez toutes les formules mathématiques en syntaxe LaTeX standard ($...$ ou \\[...\\]).\n"
                "3. Convertissez les listes en \\begin{itemize} ou \\begin{enumerate}.\n"
                "4. Convertissez les tableaux en environnements LaTeX tabular avec booktabs (\\toprule, \\midrule, \\bottomrule).\n"
                "5. Ne retournez QUE le corps du document LaTeX (sans balises markdown ```latex).\n\n"
                f"Contenu Word :\n{full_text}"
            )
            clean_body = call_gemini([prompt])
            return tex_template(clean_body)
        except Exception as e:
            print(f"Gemini DOCX fallback: {e}")

    # Fallback to local python parsing
    lines: list[str] = []
    in_list = False
    for para in document.paragraphs:
        style = para.style.name.lower() if para.style else ''
        text = para.text.strip()
        if not text:
            if in_list:
                lines.append('\\end{itemize}')
                in_list = False
            continue

        if 'heading' in style:
            if in_list:
                lines.append('\\end{itemize}')
                in_list = False
            level = 1
            try:
                level = int(style.replace('heading', '').strip())
            except ValueError:
                pass
            section_cmd = 'sub' * max(0, level - 1) + 'section'
            lines.append(f'\\{section_cmd}{{{escape_latex(text)}}}')
        elif para.style.name.startswith('List') or 'bullet' in style:
            if not in_list:
                lines.append('\\begin{itemize}')
                in_list = True
            lines.append(f'\\item {escape_latex(text)}')
        else:
            if in_list:
                lines.append('\\end{itemize}')
                in_list = False
            lines.append(f'{escape_latex(text)}\\\\')

    if in_list:
        lines.append('\\end{itemize}')

    for table in document.tables:
        col_count = len(table.columns)
        if col_count == 0:
            continue
        lines.append('\\begin{table}[h]')
        lines.append('\\centering')
        lines.append(f'\\begin{{tabular}}{{{ "l" * col_count }}}')
        lines.append('\\toprule')
        for i, row in enumerate(table.rows):
            row_cells = [escape_latex(cell.text.strip()) for cell in row.cells]
            lines.append(' & '.join(row_cells) + ' \\\\')
            if i == 0:
                lines.append('\\midrule')
        lines.append('\\bottomrule')
        lines.append('\\end{tabular}')
        lines.append('\\end{table}')

    return tex_template('\n'.join(lines))


def xlsx_to_latex(xlsx_path: Path) -> str:
    workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
    sheet_data_parts: list[str] = []

    for sheet in workbook.worksheets:
        sheet_data_parts.append(f"Feuille: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                cells = [str(cell).strip() if cell is not None else '' for cell in row]
                sheet_data_parts.append(" | ".join(cells))
        sheet_data_parts.append("")

    full_sheet_text = "\n".join(sheet_data_parts)

    if GEMINI_API_KEY and full_sheet_text.strip():
        try:
            prompt = (
                "Convertissez les données suivantes extraites d'un classeur Excel (.xlsx) en tableaux LaTeX soignés "
                "avec le package booktabs (\\toprule, \\midrule, \\bottomrule, alignement propre des colonnes, \\section* pour chaque feuille).\n"
                "Ne retournez QUE le corps du document LaTeX sans balises markdown.\n\n"
                f"Données Excel :\n{full_sheet_text}"
            )
            clean_body = call_gemini([prompt])
            return tex_template(clean_body)
        except Exception as e:
            print(f"Gemini XLSX fallback: {e}")

    # Fallback to local table creation
    body: list[str] = []
    for sheet in workbook.worksheets:
        body.append(f'\\section*{{{escape_latex(sheet.title)}}}')
        max_col = sheet.max_column or 0
        if max_col == 0:
            continue
        body.append(f'\\begin{{tabular}}{{{ "l" * max_col }}}')
        body.append('\\toprule')
        is_first = True
        for row in sheet.iter_rows(values_only=True):
            cells = [escape_latex(str(cell)) if cell is not None else '' for cell in row]
            body.append(' & '.join(cells) + ' \\\\')
            if is_first:
                body.append('\\midrule')
                is_first = False
        body.append('\\bottomrule')
        body.append('\\end{tabular}')
        body.append('')
    return tex_template('\n'.join(body))


def auto_fix_latex_syntax(broken_tex: str, compile_error: str) -> str:
    """Uses Gemini to auto-repair broken LaTeX syntax when pdflatex encounters an error."""
    if not GEMINI_API_KEY:
        return broken_tex
    try:
        prompt = (
            "Le code LaTeX suivant a échoué à la compilation pdflatex.\n"
            f"Erreur de compilation :\n{compile_error[-500:]}\n\n"
            "Corrigez les erreurs de syntaxe (par exemple caractères non échappés comme % ou &, accolades non fermées, environnement mathématique mal formé, tabular cassé).\n"
            "Retournez le document LaTeX COMPLET et corrigé, sans balises markdown."
            f"\n\nCode LaTeX à corriger :\n{broken_tex}"
        )
        fixed_text = call_gemini([prompt])
        if r'\begin{document}' not in fixed_text:
            fixed_text = tex_template(fixed_text)
        return fixed_text
    except Exception:
        return broken_tex


def find_latex_compiler() -> tuple[str, list[str]]:
    if shutil.which('tectonic'):
        return 'tectonic', ['tectonic', '{tex}', '--outdir', '{outdir}', '--keep-logs', '--quiet']
    if shutil.which('pdflatex'):
        return 'pdflatex', ['pdflatex', '-enable-installer', '-interaction=nonstopmode', '-output-directory', '{outdir}', '{tex}']
    if shutil.which('xelatex'):
        return 'xelatex', ['xelatex', '-interaction=nonstopmode', '-output-directory', '{outdir}', '{tex}']
    if shutil.which('lualatex'):
        return 'lualatex', ['lualatex', '-interaction=nonstopmode', '-output-directory', '{outdir}', '{tex}']
    raise FileNotFoundError("Aucun compilateur LaTeX (tectonic, pdflatex, xelatex, lualatex) n'a été trouvé sur le système.")


def compile_latex(tex_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / (tex_path.stem + '.pdf')
    compiler_name, cmd_template = find_latex_compiler()

    cmd = [arg.replace('{tex}', str(tex_path)).replace('{outdir}', str(output_dir)) for arg in cmd_template]

    process = subprocess.run(cmd, capture_output=True, text=True, timeout=300, cwd=output_dir)

    if process.returncode != 0 and not pdf_path.exists():
        log_file = output_dir / (tex_path.stem + '.log')
        log_content = ''
        if log_file.exists():
            log_lines = log_file.read_text(encoding='latin-1', errors='ignore').splitlines()
            log_content = '\n'.join(log_lines[-25:])
        error_msg = f"Échec de la compilation avec {compiler_name} (code {process.returncode}).\n{process.stderr}\n{log_content}"
        
        # Try auto-repair with Gemini if compilation fails
        try:
            original_tex = tex_path.read_text(encoding='utf-8', errors='ignore')
            repaired_tex = auto_fix_latex_syntax(original_tex, error_msg)
            if repaired_tex != original_tex:
                tex_path.write_text(repaired_tex, encoding='utf-8')
                retry_proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300, cwd=output_dir)
                if retry_proc.returncode == 0 and pdf_path.exists():
                    return pdf_path
        except Exception:
            pass

        raise RuntimeError(error_msg.strip())

    return pdf_path
