from pathlib import Path

fixtures = Path(__file__).resolve().parent / 'fixtures'
fixtures.mkdir(parents=True, exist_ok=True)

try:
    import docx
    doc = docx.Document()
    doc.add_heading('Exemple de document', level=1)
    doc.add_paragraph('Ceci est un test de conversion.')
    doc.save(fixtures / 'sample.docx')
except Exception as exc:
    print('docx generation failed:', exc)
    (fixtures / 'sample.docx').write_bytes(b'')

try:
    import openpyxl
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Feuille1'
    ws.append(['Col1', 'Col2', 'Col3'])
    ws.append(['A', 'B', 'C'])
    wb.save(fixtures / 'sample.xlsx')
except Exception as exc:
    print('xlsx generation failed:', exc)
    (fixtures / 'sample.xlsx').write_bytes(b'')

pdf_content = b"""%PDF-1.1
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Count 1/Kids [3 0 R]>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox [0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length 44>>stream
BT /F1 24 Tf 100 700 Td (Hello PDF) Tj ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f 
0000000010 00000 n 
0000000060 00000 n 
0000000115 00000 n 
0000000240 00000 n 
0000000323 00000 n 
trailer<</Root 1 0 R/Size 6>>
startxref
384
%%EOF
"""
(fixtures / 'sample.pdf').write_bytes(pdf_content)
print('fixtures generated')
